#!/usr/bin/env python
# coding: utf-8

# In[111]:


"""Convert a csv file to a table in a Microsoft Word document."""
import os
import numpy as np
import pandas as pd
import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.shared import Inches


# In[170]:


def format_number(x):
    """Check type of `x` and format it appropriately."""
    if isinstance(x, float):
        return f"{x:.3f}"
    elif isinstance(x, (int, np.integer)):
        return f"{x:,}"
    else:
        return str(x)

def add_asterisks_string(s, n):
    return s + '*' * n

def add_asterisks_column(df, col_number, n):
    f = lambda s: add_asterisks_string(s, n)
    df.iloc[:, col_number] = df.iloc[:, col_number].apply(f)


# In[180]:


filename = "excel2.xlsx" # Change this. Should be the name of the file without file extension
directory = "private"
df = pd.read_excel(os.path.join(directory, filename), header=[0,1])
df = df.applymap(format_number) # Format the numbers
add_asterisks_column(df, 1, 3)
add_asterisks_column(df, 6, 3)
df


# In[181]:


df.shape
df.iloc[15, 1]


# In[182]:


# Create an empty document
doc = docx.Document()

# Adjust margins
sections = doc.sections
for section in sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)

ncolumns = df.shape[-1] # Number of columns
nlevels = df.columns.nlevels # Number of header rows
nrows = df.shape[0] # Number of rows

table = doc.add_table(nlevels + nrows, ncolumns)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# add the header rows.
for i in range(nlevels):
    for j in range(ncolumns):
        text = df.columns[j][i]
        # Skip unnamed cells
        if text[:8] == "Unnamed:":
            continue
        
        if j==0 or text != table.cell(i, j-1).text:
            table.cell(i, j).text = text
            table.cell(i, j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        else:
            table.cell(i, j-1).merge(table.cell(i, j))

# add the rest of the data frame
for i in range(nrows):
    for j in range(ncolumns):
        table.cell(i + nlevels, j).text = df.iloc[i, j]
        table.cell(i + nlevels, j).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER


# In[183]:


# Save the document
root = filename.split('.')[0]
doc.save(os.path.join(directory, f"{root}.docx"))


# In[ ]:




